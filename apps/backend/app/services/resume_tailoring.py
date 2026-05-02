from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.candidate import Candidate
from app.models.job import JobNormalized
from app.models.resume_tailoring import ResumeTailoringDraft

logger = logging.getLogger(__name__)

RESUME_DIR = Path(os.getenv("RESUME_STORAGE_PATH", "/app/resumes"))
_STOP_WORDS = {
    "and", "are", "for", "from", "have", "with", "this", "that", "will", "your",
    "you", "the", "job", "role", "team", "work", "using", "used", "years", "must",
    "plus", "nice", "required", "preferred", "experience", "skills", "ability",
}


def create_tailoring_draft(
    db: Session,
    *,
    candidate: Candidate,
    job: JobNormalized,
    match_id: int | None,
    recruiter_context: str | None,
    created_by_user_id: int | None,
) -> ResumeTailoringDraft:
    suggestions, gaps = build_tailoring_suggestions(candidate, job, recruiter_context)
    draft = ResumeTailoringDraft(
        candidate_id=candidate.id,
        job_id=job.id,
        match_id=match_id,
        created_by_user_id=created_by_user_id,
        recruiter_context=(recruiter_context or "").strip() or None,
        suggested_edits=suggestions,
        skill_gaps=gaps,
        approved_edits=[],
        confirmed_skills=[],
        status="draft",
    )
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return draft


def build_tailoring_suggestions(
    candidate: Candidate,
    job: JobNormalized,
    recruiter_context: str | None,
) -> tuple[list[dict], list[dict]]:
    requested_skills = _requested_skills(job, recruiter_context)
    evidence_terms = _candidate_evidence_terms(candidate)
    supported = [skill for skill in requested_skills if _has_evidence(skill, evidence_terms)]
    gaps = [
        {
            "skill": skill,
            "reason": "This skill was requested by the JD or recruiter context but was not found in the stored resume/profile.",
            "source": "job_or_recruiter_context",
        }
        for skill in requested_skills
        if skill not in supported
    ]

    ai_suggestions = _build_ai_suggestions(candidate, job, recruiter_context, supported)
    if ai_suggestions:
        return ai_suggestions, gaps

    suggestions: list[dict] = []
    for skill in supported[:8]:
        suggestions.append(
            {
                "id": f"sug_{uuid4().hex[:10]}",
                "section": "Targeted Highlights",
                "text": f"Applied {skill} experience to support {job.title} responsibilities in healthcare and business operations environments.",
                "skill_tags": [skill],
                "evidence": f"Found supporting evidence for {skill} in the candidate resume/profile.",
                "status": "supported",
            }
        )
    if supported:
        suggestions.insert(
            0,
            {
                "id": f"sug_{uuid4().hex[:10]}",
                "section": "Professional Summary",
                "text": (
                    f"Healthcare IT professional aligned to {job.title} requirements, with relevant strengths in "
                    f"{', '.join(supported[:5])}."
                ),
                "skill_tags": supported[:5],
                "evidence": "Built from skills already present in the candidate resume/profile.",
                "status": "supported",
            },
        )
    return suggestions, gaps


def generate_tailored_docx(
    db: Session,
    *,
    draft: ResumeTailoringDraft,
    candidate: Candidate,
    job: JobNormalized,
    approved_suggestion_ids: list[str],
    confirmed_skills: list[str],
) -> Path:
    source = _source_docx_path(candidate)
    approved = [
        item for item in draft.suggested_edits
        if item.get("id") in set(approved_suggestion_ids) and item.get("status") == "supported"
    ]
    confirmed_lookup = {skill.strip().lower() for skill in confirmed_skills if skill.strip()}
    for gap in draft.skill_gaps:
        skill = str(gap.get("skill", "")).strip()
        if skill and skill.lower() in confirmed_lookup:
            approved.append(
                {
                    "id": f"confirmed_{uuid4().hex[:10]}",
                    "section": "Targeted Highlights",
                    "text": f"Added confirmed experience with {skill} for {job.title} requirements.",
                    "skill_tags": [skill],
                    "evidence": "Employee confirmed this skill is accurate for the candidate.",
                    "status": "confirmed",
                }
            )

    if not approved:
        raise ValueError("Select at least one supported edit or confirm one requested skill before downloading.")

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate tailored Word resumes.") from exc

    RESUME_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(str(source))
    document.add_page_break()
    document.add_heading("Targeted Resume Enhancements", level=2)
    document.add_paragraph(f"Role focus: {job.title} at {job.company}")
    if job.apply_url or job.canonical_apply_url:
        document.add_paragraph(f"Apply link: {job.apply_url or job.canonical_apply_url}")

    grouped: dict[str, list[dict]] = {}
    for item in approved:
        grouped.setdefault(str(item.get("section") or "Targeted Highlights"), []).append(item)

    for section, items in grouped.items():
        document.add_heading(section, level=3)
        for item in items:
            document.add_paragraph(str(item.get("text", "")).strip(), style="List Bullet")

    filename = f"candidate_{candidate.id}_tailored_job_{job.id}_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.docx"
    output_path = RESUME_DIR / filename
    document.save(str(output_path))

    draft.approved_edits = approved
    draft.confirmed_skills = sorted({skill.strip() for skill in confirmed_skills if skill.strip()})
    draft.generated_filename = filename
    draft.status = "generated"
    db.commit()
    db.refresh(draft)
    return output_path


def _source_docx_path(candidate: Candidate) -> Path:
    if not candidate.resume_filename:
        raise ValueError("No master resume is uploaded for this candidate.")
    if Path(candidate.resume_filename).suffix.lower() != ".docx":
        raise ValueError("Tailored Word download requires the candidate's master resume to be a DOCX file.")
    path = RESUME_DIR / candidate.resume_filename
    if not path.exists():
        raise ValueError("Master resume file was not found on disk.")
    return path


def _requested_skills(job: JobNormalized, recruiter_context: str | None) -> list[str]:
    skills: list[str] = []
    skills.extend(job.keywords_extracted or [])
    skills.extend(_extract_skill_phrases(recruiter_context or ""))
    skills.extend(_extract_skill_phrases(job.description or ""))
    return _dedupe_preserve_order([_canonical_skill(skill) for skill in skills if _canonical_skill(skill)])[:30]


def _extract_skill_phrases(text: str) -> list[str]:
    raw_tokens = re.findall(r"\b[A-Za-z][A-Za-z0-9+#.\-/]{1,30}\b", text)
    return [
        token.upper() if token.lower() in {"sql", "fhir", "hl7", "edi", "api", "etl", "qa", "uat"} else token
        for token in raw_tokens
        if len(token) > 2 and token.lower() not in _STOP_WORDS
    ]


def _candidate_evidence_terms(candidate: Candidate) -> set[str]:
    text = " ".join(
        [
            candidate.resume_text or "",
            " ".join(skill.skill_name for skill in candidate.skills),
            " ".join(candidate.preference.must_have_keywords if candidate.preference else []),
            " ".join(candidate.preference.domain_expertise if candidate.preference else []),
        ]
    )
    return {_canonical_skill(token).lower() for token in _extract_skill_phrases(text)}


def _has_evidence(skill: str, evidence_terms: set[str]) -> bool:
    normalized = skill.lower()
    return normalized in evidence_terms or any(normalized in term or term in normalized for term in evidence_terms)


def _canonical_skill(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip(" ,.;:()[]{}")).strip()


def _dedupe_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        key = value.lower()
        if key and key not in seen:
            seen.add(key)
            result.append(value)
    return result


def _build_ai_suggestions(
    candidate: Candidate,
    job: JobNormalized,
    recruiter_context: str | None,
    supported_skills: list[str],
) -> list[dict]:
    if not settings.ai_scoring_enabled or not supported_skills:
        return []
    try:
        raw = _call_ai(_prompt(candidate, job, recruiter_context, supported_skills))
        data = json.loads(raw.strip())
        items = data.get("suggestions", [])
        suggestions: list[dict] = []
        for item in items[:10]:
            text = str(item.get("text", "")).strip()
            if not text:
                continue
            suggestions.append(
                {
                    "id": f"sug_{uuid4().hex[:10]}",
                    "section": str(item.get("section") or "Targeted Highlights"),
                    "text": text,
                    "skill_tags": [str(skill) for skill in item.get("skill_tags", []) if str(skill) in supported_skills],
                    "evidence": str(item.get("evidence") or "Supported by candidate resume/profile."),
                    "status": "supported",
                }
            )
        return suggestions
    except Exception as exc:
        logger.warning("resume_tailoring.ai_failed: %s", exc)
        return []


def _prompt(candidate: Candidate, job: JobNormalized, recruiter_context: str | None, supported_skills: list[str]) -> str:
    return f"""You help staffing employees tailor resumes truthfully.

Only write resume bullets using these supported skills: {", ".join(supported_skills[:20])}.
Do not invent employers, certifications, tools, metrics, or years.

Candidate resume excerpt:
{(candidate.resume_text or "")[:2500]}

Job:
{job.title} at {job.company}
{(job.description or "")[:2500]}

Recruiter context:
{(recruiter_context or "")[:1200]}

Return exactly JSON:
{{"suggestions":[{{"section":"Professional Summary|Skills|Targeted Highlights","text":"resume-ready sentence or bullet","skill_tags":["skill"],"evidence":"short evidence note"}}]}}"""


def _call_ai(prompt: str) -> str:
    provider = settings.ai_provider.lower()
    if provider == "anthropic":
        import anthropic  # type: ignore

        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        msg = client.messages.create(
            model=settings.ai_model,
            max_tokens=900,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text

    import openai  # type: ignore

    client = openai.OpenAI(api_key=settings.openai_api_key)
    resp = client.chat.completions.create(
        model=settings.ai_model,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=900,
        temperature=0.2,
    )
    return resp.choices[0].message.content or ""
