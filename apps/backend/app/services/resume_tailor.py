"""Resume tailoring service.

Given a candidate's master DOCX resume and a job description, calls OpenAI to
generate targeted edits and applies them using python-docx.  The master file is
NEVER modified — output goes to UPLOAD_DIR/tailored/.

Design principles:
- Employee notes are treated as PRIMARY instructions.  GPT only does what the
  employee explicitly asks — no unsolicited rewrites.
- New paragraphs copy paragraph & run formatting from their insertion point,
  preserving bullet styles, indentation, and fonts.
- Changes are insert_after operations identified by paragraph index so there is
  no ambiguous text-search.
"""
from __future__ import annotations

import copy
import io
import json
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from app.models.candidate import Candidate, CandidateSkill
from app.models.tailored_resume import TailoredResume

_RESUME_DIR = Path(os.getenv("RESUME_STORAGE_PATH", "/app/resumes"))
_TAILORED_DIR = _RESUME_DIR / "tailored"
_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ── DOCX utilities ────────────────────────────────────────────────────────────

def _extract_docx_text(path: Path) -> str:
    import mammoth  # type: ignore
    result = mammoth.extract_raw_text({"path": str(path)})
    return result.value or ""


def _extract_resume_text_from_bytes(content: bytes, suffix: str) -> str:
    """Best-effort plain text extraction for suggestion-only fallback modes."""
    suffix = suffix.lower()
    if suffix == ".pdf":
        import pypdf  # type: ignore
        reader = pypdf.PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix in {".doc", ".docx"}:
        import mammoth  # type: ignore
        result = mammoth.extract_raw_text({"bytes": content})
        return result.value or ""
    return content.decode("utf-8", errors="ignore")


def _candidate_resume_bytes(candidate: Candidate | None, path: Path) -> bytes | None:
    if candidate and candidate.resume_bytes:
        return candidate.resume_bytes
    if path.exists():
        return path.read_bytes()
    return None


def _save_suggestions_record(
    record: TailoredResume,
    suggestions: list[str],
    message: str,
    db: Session,
) -> tuple[TailoredResume, list[str]]:
    if not suggestions:
        record.status = "error"
        record.error_message = "No truthful copy-paste lines could be generated from this resume and job description."
        db.commit()
        return record, []
    record.status = "suggestions_ready"
    record.suggested_lines = "\n".join(suggestions)
    record.error_message = message
    db.commit()
    db.refresh(record)
    return record, []


def _restore_resume_cache(candidate: Candidate | None, path: Path) -> bool:
    content = _candidate_resume_bytes(candidate, path)
    if not content:
        return False
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(content)
    except Exception:
        return path.exists()
    return True


def _get_doc_paragraphs(path: Path) -> list[dict[str, Any]]:
    """Return a numbered list of non-empty paragraphs for GPT context."""
    import docx  # python-docx
    doc = docx.Document(str(path))
    paras = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            paras.append({
                "index": i,
                "style": p.style.name if p.style else "Normal",
                "text": text[:300],
            })
    return paras


def _make_paragraph_element(doc: Any, text: str, ref_para: Any) -> Any:
    """Build a new <w:p> element that inherits paragraph + run formatting from ref_para."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    new_p = OxmlElement("w:p")

    # Copy paragraph properties (indentation, bullet style, spacing, etc.)
    ref_pPr = ref_para._element.find(qn("w:pPr"))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))

    # Build run
    new_r = OxmlElement("w:r")

    # Copy run properties (font, size, bold, etc.) from first run of ref_para
    if ref_para.runs:
        ref_rPr = ref_para.runs[0]._element.find(qn("w:rPr"))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))

    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)

    return new_p


def _apply_changes_to_docx(source_path: Path, changes: list[dict[str, Any]]) -> bytes:
    """Apply insert_after changes to an in-memory DOCX copy.

    Each change must have:
      action          : "insert_after"
      paragraph_index : int  — insert new para after para at this index
      content         : str  — text for the new paragraph
    """
    import docx  # python-docx

    doc = docx.Document(str(source_path))

    # Snapshot paragraph list BEFORE any mutation so indices stay stable
    para_snapshot = list(doc.paragraphs)

    # Process in REVERSE index order so earlier inserts don't shift later ones
    inserts = sorted(
        [
            c for c in changes
            if c.get("action") == "insert_after"
            and isinstance(c.get("paragraph_index"), int)
            and c.get("content", "").strip()
        ],
        key=lambda x: x["paragraph_index"],
        reverse=True,
    )

    for change in inserts:
        idx = change["paragraph_index"]
        content = change["content"].strip()

        if idx < 0 or idx >= len(para_snapshot):
            continue

        ref_para = para_snapshot[idx]
        new_elem = _make_paragraph_element(doc, content, ref_para)
        ref_para._element.addnext(new_elem)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── OpenAI call ───────────────────────────────────────────────────────────────

def _call_openai(
    doc_paragraphs: list[dict[str, Any]],
    jd_text: str,
    notes: str | None,
) -> list[dict[str, Any]]:
    """Return a list of insert_after change dicts from GPT-4o.

    If employee notes are provided, they are treated as the ONLY instruction.
    GPT is not allowed to make unsolicited changes.
    """
    from openai import OpenAI
    from app.core.config import settings

    if not settings.openai_api_key:
        raise RuntimeError("OPENAI_API_KEY is not configured. Add it to the backend environment and try again.")

    client = OpenAI(api_key=settings.openai_api_key)

    # Build the paragraph context string
    para_lines = "\n".join(
        f"[{p['index']}] ({p['style']}) {p['text']}"
        for p in doc_paragraphs
    )

    if notes and notes.strip():
        instructions_block = f"""
═══ EMPLOYEE INSTRUCTIONS (follow EXACTLY and ONLY these) ═══
{notes.strip()}
═══════════════════════════════════════════════════════════════

CRITICAL RULES:
1. ONLY perform the action described in the employee instructions above.
2. Do NOT change, rewrite, or touch anything not explicitly mentioned.
3. Do NOT add a summary, do NOT change skills, do NOT reformat.
4. If the instruction says "last two projects" or "last two jobs": find the
   two most recent job title/company lines (usually highest-indexed job
   entries in the Experience section) and insert new bullet points
   IMMEDIATELY AFTER those lines.
5. Match the bullet style already used in the document (copy from nearby
   bullets — same indentation, same character).
6. Each bullet should be a single, specific, professional responsibility.
7. Keep changes minimal — only what was asked for.
"""
    else:
        instructions_block = """
Make minimal, targeted edits to better align the resume with the job description.
Add 2-3 concise bullets/lines near the most recent project or latest experience only.
Only add job-description elements that are already supported by the resume context.
Do not reformat or rewrite whole sections.
"""

    system_prompt = f"""You are a precise resume editor. Your output is ONLY a JSON array of insertions.

Each insertion is:
{{
  "action": "insert_after",
  "paragraph_index": <int>,   // insert new paragraph after this paragraph index
  "content": "<text>"         // the exact text for the new bullet/line
}}

{instructions_block}

RULES:
- Return ONLY a valid JSON array. No markdown fences. No prose. No explanation.
- Empty array [] if no changes are needed.
- paragraph_index must be an integer from the paragraph list provided.
- Do NOT invent companies, degrees, or certifications that aren't implied by the existing resume.
"""

    user_msg = (
        f"RESUME PARAGRAPHS (index → style → text):\n{para_lines[:5000]}\n\n"
        f"JOB DESCRIPTION:\n{jd_text[:2000]}"
    )

    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_msg},
        ],
        temperature=0.2,
        max_tokens=1500,
    )

    raw = (resp.choices[0].message.content or "").strip()

    # Strip markdown code fences if GPT wraps the JSON
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    try:
        changes = json.loads(raw)
        if not isinstance(changes, list):
            changes = []
    except (json.JSONDecodeError, ValueError):
        changes = []

    return changes[:3]


def _call_openai_suggestions(
    resume_text: str,
    jd_text: str,
    notes: str | None,
) -> list[str]:
    """Return 2-3 truthful copy-paste lines when the source resume is not editable DOCX."""
    from openai import OpenAI
    from app.core.config import settings

    if not settings.openai_api_key:
        raise RuntimeError("OPENAI_API_KEY is not configured. Add it to the backend environment and try again.")

    client = OpenAI(api_key=settings.openai_api_key)
    extra_notes = f"\nEMPLOYEE NOTES:\n{notes.strip()}\n" if notes and notes.strip() else ""
    system_prompt = """You write resume bullets for copy/paste into the candidate's latest project.

Return ONLY a JSON array of 2-3 strings. No markdown, no prose.
Rules:
- Each string must be one concise resume bullet or project line.
- Use only experience, tools, domains, and facts supported by the resume text.
- Target missing job-description elements only when they are truthful for this candidate.
- Do not invent employers, degrees, certifications, metrics, tools, or responsibilities.
- If there is not enough support in the resume, return the best truthful transferable lines.
"""
    user_msg = (
        f"RESUME TEXT:\n{resume_text[:5000]}\n\n"
        f"JOB DESCRIPTION:\n{jd_text[:2500]}\n"
        f"{extra_notes}"
    )
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.2,
        max_tokens=900,
    )
    raw = (resp.choices[0].message.content or "").strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()
    try:
        data = json.loads(raw)
    except (json.JSONDecodeError, ValueError):
        data = []
    if not isinstance(data, list):
        return []
    return [str(line).strip() for line in data if str(line).strip()][:3]


# ── Skill flagging ─────────────────────────────────────────────────────────────

def _get_candidate_skill_names(db: Session, candidate_id: int) -> set[str]:
    rows = db.query(CandidateSkill).filter(CandidateSkill.candidate_id == candidate_id).all()
    return {r.skill_name.lower() for r in rows}


def _flag_unknown_skills(
    changes: list[dict[str, Any]],
    known_skills: set[str],
    resume_text: str,
) -> list[str]:
    """Return likely tools/skills in inserted content that are not evidenced."""
    resume_lower = resume_text.lower()
    common_words = {
        "and", "for", "with", "from", "into", "using", "used", "led", "the",
        "that", "this", "across", "business", "project", "projects", "teams",
        "requirements", "process", "data", "system", "systems", "analysis",
        "stakeholders", "workflows", "solutions", "support", "improve",
        "improved", "managed", "created", "developed", "designed",
    }
    candidates: list[str] = []
    for change in changes:
        content = str(change.get("content", ""))
        # Likely skill/tool terms: acronyms, terms with symbols, or title-cased
        # multiword fragments. Avoid flagging every ordinary resume word.
        tokens = re.findall(r"\b[A-Za-z][A-Za-z0-9+#./-]{2,}\b", content)
        for token in tokens:
            normalized = token.strip(",.;()").lower()
            if normalized in common_words or normalized in known_skills or normalized in resume_lower:
                continue
            if token.isupper() or any(ch in token for ch in "+#./-"):
                candidates.append(token)
    return list(dict.fromkeys(candidates))


# ── Main entry point ──────────────────────────────────────────────────────────

def tailor_resume(
    candidate_id: int,
    job_id: int,
    master_resume_path: str,
    jd_text: str,
    notes: str | None,
    db: Session,
    created_by_employee_id: int | None = None,
    match_id: int | None = None,
    confirm_flagged_skills: list[str] | None = None,
) -> tuple[TailoredResume, list[str]]:
    """Tailor a candidate's master DOCX resume for a specific job.

    Returns (TailoredResume record, flagged_skills list).
    If flagged_skills is non-empty and confirm_flagged_skills is None, the
    record stays at status='pending' for the employee to confirm before proceeding.
    """
    path = Path(master_resume_path)
    suffix = path.suffix.lower()
    candidate = db.get(Candidate, candidate_id)

    # Create the DB record first — gives us an ID to return in all error paths
    record = TailoredResume(
        candidate_id=candidate_id,
        job_id=job_id,
        match_id=match_id,
        created_by_employee_id=created_by_employee_id,
        status="pending",
        notes=notes,
    )
    db.add(record)
    db.commit()
    db.refresh(record)

    # ── Validate master resume ───────────────────────────────────────────────
    if not candidate or not candidate.resume_filename:
        record.status = "error"
        record.error_message = "Candidate has no resume on file. Please upload a resume first."
        db.commit()
        return record, []

    source_bytes = _candidate_resume_bytes(candidate, path)
    if not source_bytes:
        resume_text = (candidate.resume_text or "").strip()
        if resume_text:
            record.status = "processing"
            db.commit()
            try:
                suggestions = _call_openai_suggestions(resume_text, jd_text, notes)
            except Exception as exc:
                record.status = "error"
                record.error_message = f"AI service error: {exc}"
                db.commit()
                return record, []
            return _save_suggestions_record(
                record,
                suggestions,
                (
                    "The original resume file is missing from storage, so a tailored download cannot be created. "
                    "Copy these lines into the latest project, or re-upload the resume once to enable direct DOCX downloads."
                ),
                db,
            )
        record.status = "error"
        record.error_message = (
            "Resume file not found in storage. Please re-upload the candidate's resume once so the system can save it."
        )
        db.commit()
        return record, []

    if suffix != ".docx":
        record.status = "processing"
        db.commit()
        try:
            resume_text = (candidate.resume_text or "").strip() or _extract_resume_text_from_bytes(source_bytes, suffix)
        except Exception as exc:
            record.status = "error"
            record.error_message = f"Could not read resume text for suggestions: {exc}"
            db.commit()
            return record, []
        if not resume_text.strip():
            record.status = "error"
            record.error_message = "The resume appears to be empty or unreadable. Please re-upload a readable resume."
            db.commit()
            return record, []
        try:
            suggestions = _call_openai_suggestions(resume_text, jd_text, notes)
        except Exception as exc:
            record.status = "error"
            record.error_message = f"AI service error: {exc}"
            db.commit()
            return record, []
        return _save_suggestions_record(
            record,
            suggestions,
            (
                "This resume format cannot be safely edited automatically. "
                "Copy these suggested lines into the latest project, or upload DOCX for direct download."
            ),
            db,
        )

    if not path.exists() and not _restore_resume_cache(candidate, path):
        record.status = "error"
        record.error_message = "Resume file not found. Please re-upload the candidate's .docx resume."
        db.commit()
        return record, []

    # ── Extract paragraph structure ──────────────────────────────────────────
    try:
        doc_paragraphs = _get_doc_paragraphs(path)
    except Exception as exc:
        record.status = "error"
        record.error_message = f"Could not read the DOCX file: {exc}"
        db.commit()
        return record, []

    if not doc_paragraphs:
        record.status = "error"
        record.error_message = "The resume appears to be empty or unreadable."
        db.commit()
        return record, []

    # ── Call OpenAI ──────────────────────────────────────────────────────────
    record.status = "processing"
    db.commit()

    try:
        changes = _call_openai(doc_paragraphs, jd_text, notes)
    except Exception as exc:
        record.status = "error"
        record.error_message = f"AI service error: {exc}"
        db.commit()
        return record, []

    if not changes:
        # GPT found nothing to change — still mark ready with original
        # (copy master to tailored dir unchanged)
        ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        out_filename = f"tailored_{candidate_id}_{job_id}_{ts}.docx"
        out_path = _TAILORED_DIR / out_filename
        import shutil
        try:
            _TAILORED_DIR.mkdir(parents=True, exist_ok=True)
            shutil.copy2(path, out_path)
        except Exception:
            pass
        record.file_bytes = source_bytes
        record.content_type = _DOCX_CONTENT_TYPE
        record.status = "ready"
        record.filename = out_filename
        record.error_message = "No changes were needed — the resume already aligns well."
        db.commit()
        db.refresh(record)
        return record, []

    # ── Skill flagging ───────────────────────────────────────────────────────
    known_skills = _get_candidate_skill_names(db, candidate_id)
    flagged_skills: list[str] = []
    if known_skills:
        resume_text = (candidate.resume_text or "").strip() or _extract_docx_text(path)
        flagged_skills = _flag_unknown_skills(changes, known_skills, resume_text)

    if flagged_skills and confirm_flagged_skills is None:
        record.status = "pending"
        db.commit()
        return record, flagged_skills

    # Filter out unconfirmed flagged skills
    confirmed_set = {s.lower() for s in (confirm_flagged_skills or [])}
    if flagged_skills:
        def _is_confirmed(change: dict[str, Any]) -> bool:
            words = {w.strip(",.;()").lower() for w in change.get("content", "").split()}
            unconfirmed = [s for s in flagged_skills if s not in confirmed_set and s in words]
            return len(unconfirmed) == 0
        changes = [c for c in changes if _is_confirmed(c)]

    # ── Apply changes to DOCX ────────────────────────────────────────────────
    try:
        docx_bytes = _apply_changes_to_docx(path, changes)
    except Exception as exc:
        record.status = "error"
        record.error_message = f"Failed to apply changes to document: {exc}"
        db.commit()
        return record, []

    # ── Save tailored DOCX ───────────────────────────────────────────────────
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    out_filename = f"tailored_{candidate_id}_{job_id}_{ts}.docx"
    out_path = _TAILORED_DIR / out_filename
    try:
        _TAILORED_DIR.mkdir(parents=True, exist_ok=True)
        out_path.write_bytes(docx_bytes)
    except Exception:
        pass

    record.status = "ready"
    record.filename = out_filename
    record.file_bytes = docx_bytes
    record.content_type = _DOCX_CONTENT_TYPE
    db.commit()
    db.refresh(record)
    return record, []
